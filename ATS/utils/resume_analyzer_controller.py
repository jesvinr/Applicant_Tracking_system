import re
import requests
from PyPDF2 import PdfReader
from docx import Document
from datetime import datetime
import pandas as pd
from io import BytesIO

class ResumeAnalyzer:
    def __init__(self):
        # Document type indicators
        self.document_types = {
            'resume': [
                'experience', 'education', 'skills', 'work', 'project', 'objective',
                'summary', 'employment', 'qualification', 'achievements'
            ],
            'marksheet': [
                'grade', 'marks', 'score', 'semester', 'cgpa', 'sgpa', 'examination',
                'result', 'academic year', 'percentage'
            ],
            'certificate': [
                'certificate', 'certification', 'awarded', 'completed', 'achievement',
                'training', 'course completion', 'qualified'
            ],
            'id_card': [
                'id card', 'identity', 'student id', 'employee id', 'valid until',
                'date of issue', 'identification'
            ]
        }
        
    def detect_document_type(self, text):
        text = text.lower()
        scores = {}
        
        # Calculate score for each document type
        for doc_type, keywords in self.document_types.items():
            matches = sum(1 for keyword in keywords if keyword in text)
            density = matches / len(keywords)
            frequency = matches / (len(text.split()) + 1)  # Add 1 to avoid division by zero
            scores[doc_type] = (density * 0.7) + (frequency * 0.3)
        
        # Get the highest scoring document type
        best_match = max(scores.items(), key=lambda x: x[1])
        
        # Only return a document type if the score is significant
        return best_match[0] if best_match[1] > 0.15 else 'unknown'
        
    def calculate_keyword_match(self, resume_text, required_skills):
        resume_text = resume_text.lower()
        found_skills = []
        missing_skills = []
        
        for skill in required_skills:
            skill_lower = skill.lower()
            # Check for exact match
            if skill_lower in resume_text:
                found_skills.append(skill)
            # Check for partial matches (e.g., "Python" in "Python programming")
            elif any(skill_lower in phrase for phrase in resume_text.split('.')):
                found_skills.append(skill)
            else:
                missing_skills.append(skill)
                
        match_score = (len(found_skills) / len(required_skills)) * 100 if required_skills else 0
        
        return {
            'score': match_score,
            'found_skills': found_skills,
            'missing_skills': missing_skills
        }
        
    def check_resume_sections(self, text):
        text = text.lower()
        essential_sections = {
            'contact': ['email', 'phone', 'address', 'linkedin'],
            'education': ['education', 'university', 'college', 'degree', 'academic'],
            'experience': ['experience', 'work', 'employment', 'job', 'internship'],
            'skills': ['skills', 'technologies', 'tools', 'proficiencies', 'expertise']
        }
        
        section_scores = {}
        for section, keywords in essential_sections.items():
            found = sum(1 for keyword in keywords if keyword in text)
            section_scores[section] = min(25, (found / len(keywords)) * 25)
            
        return sum(section_scores.values())
        
    def check_formatting(self, text):
        lines = text.split('\n')
        score = 100
        deductions = []
        
        # Check for minimum content
        if len(text) < 300:
            score -= 30
            deductions.append("Resume is too short")
            
        # Check for section headers
        if not any(line.isupper() for line in lines):
            score -= 20
            deductions.append("No clear section headers found")
            
        # Check for bullet points
        if not any(line.strip().startswith(('•', '-', '*', '→')) for line in lines):
            score -= 20
            deductions.append("No bullet points found for listing details")
            
        # Check for consistent spacing
        if any(len(line.strip()) == 0 and len(next_line.strip()) == 0 
               for line, next_line in zip(lines[:-1], lines[1:])):
            score -= 15
            deductions.append("Inconsistent spacing between sections")
            
        # Check for contact information format
        contact_patterns = [
            r'\b[\w\.-]+@[\w\.-]+\.\w+\b',  # email
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',  # phone
            r'linkedin\.com/\w+',  # LinkedIn
        ]
        if not any(re.search(pattern, text) for pattern in contact_patterns):
            score -= 15
            deductions.append("Missing or improperly formatted contact information")
            
        return max(0, score), deductions
        
    def extract_text_from_pdf(self, file):
        try:
            import PyPDF2
            import io
            
            # Create a PDF reader object
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
            
            # Extract text from all pages
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
                
            return text
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
            
    def extract_text_from_docx(self, docx_file):
        """Extract text from a DOCX file"""
        try:
            from docx import Document
            doc = Document(docx_file)
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            return '\n'.join(full_text)
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX file: {str(e)}")

    def extract_personal_info(self, text):
        """Extract personal information from resume text"""
        # Basic patterns for personal info
        email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
        print(text)
        # email_pattern = r'([a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$)'
        phone_pattern = r'(\+\d{1,3}[-.]?)?\s*\(?\d{3}\)?[-.]?\s*\d{3}[-.]?\s*\d{4}'
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        github_pattern = r'github\.com/[\w-]+'
        
        # Extract information
        email = re.search(email_pattern, text)
        phone = re.search(phone_pattern, text)
        linkedin = re.search(linkedin_pattern, text)
        github = re.search(github_pattern, text)
        
        # Get the first line as name (basic assumption)
        name = text.split('\n')[0].strip()

        email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"

        matches = re.findall(email_pattern, str(email))  # Extracts all emails in the text

        # print(matches)  

        # print(name, email)
        # print(self.extractUniqueWordsAndSentences(text=text))
        return {
            'name': name if len(name) > 0 else 'Unknown',
            'email': email.group(0) if email else 'Unknown',
            'phone': phone.group(0) if phone else '',
            'linkedin': linkedin.group(0) if linkedin else '',
            'github': github.group(0) if github else '',
            'portfolio': ''  # Can be enhanced later
        }

    def extract_resume_details_using_LLM(self,resume_text):
        payload = {
            "model": "gemma2-9b-it",  # Change to the model you prefer
            "messages": [
                {"role": "system", "content": """YOU ARE AN AI-POWERED ATS (APPLICANT TRACKING SYSTEM) RESUME EVALUATOR, DESIGNED TO DELIVER DETAILED, 
                    PROFESSIONAL ANALYSIS OF A USER’S RESUME. YOU WILL EVALUATE THE RESUME AGAINST INDUSTRY STANDARDS FOR ATS COMPATIBILITY AND SCORING, 
                    THEN PROVIDE A COMPREHENSIVE BREAKDOWN OF YOUR FINDINGS.
                 You should provide details in given format only."""},
                {"role": "user", "content": f"""Extract the following details from the resume: 
                    - Education
                    - Experience
                    - Projects
                    Format =>
                    Education: 
                    <Education details in one line(for each education) >
                    Experience:
                    <Experience details => add from and to date and breif summary in one line>
                    Experience Dates:
                    <from date - to date, in format described in point 1 in Notes below, 
                        if current or present is give in input change it to current month, year>
                        Example: jan/2024 - jun/2024 or sep/2024 - 03/2025 or 01/2023 - 03/2023
                        if Present or Current is given [i.e jan/2023 - Current] convert it to current month i.e jan/2023 - mar/2025
                        here (mar/2025 or 03/2025) is current month.
                    Projects:
                    <project details like project name  and summary>
                    Note <for you to understand (output must be in below format), dont print it for user>: 
                    1) here every date must be in format of "<month first 3 character>/<4 digit of year>". Eg: mar/2024, jan/2020. 
                    1.1) If user not given in above format, change to above format (point 1 in notes) and display it in output [must].
                    2) no more or less details should be added other than whats asked.
                    3) for each point use "-" before it. Don't use any symbol before section heading. [no - or # or *] 
                        i.e (Education, Experience, Projects) (in a same line).
                    4) for experince dates section alone use number points for each experince. 
                    {resume_text}"""} 
            ],
            "temperature":0.1
        }
        API_KEY = "<YOUR_API_KEY>"  # Replace with your actual API key

        URL = "https://api.groq.com/openai/v1/chat/completions"
        HEADERS = {"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"}  # Replace with your API Key
        
        response = requests.post(URL, headers=HEADERS, json=payload)
        response= response.json()
        # print("Feedback:", response.get("choices", [{}])[0].get("message", {}).get("content", "No response received."))
        return response.get("choices", [{}])[0].get("message", {}).get("content", "No response received.")

    def extract_experience_dates_from_section(self,text):
        """
        Extracts experience date ranges from the 'Experience Dates' section and returns them as a list.
        """ 
        # Find the individual date ranges
        # pattern1 = r'\d\. (\w{3}|\d{2}/\d{4} - \w{3}|\d{2}/\d{4})'
        # pattern2 = r'\d\. (\w{3}|\d{2}/\d{4} - [A-Za-z]+)'
        # date_range_pattern = f'({pattern1})|({pattern2})'
        pattern1 = r'\d+\.\s*(\d{2}/\d{4})\s*-\s*(\d{2}/\d{4}|current)'
        pattern2 = r'\d+\.\s*([A-Za-z]{3}/\d{4})\s*-\s*([A-Za-z]{3}/\d{4}|current)'
        # date_range_pattern = f'{pattern1}|{pattern2}' #r'\d\. (\w{3}|\d{2}/\d{4} - [A-Za-z]+|[0-9/]{7})'
        #r'\d+\. \s*(\d{2}/\d{4}\s*-\s*\w{3}|\d{2}/\d{4}|current)' #f'({pattern1})|({pattern2})'

        date_range_pattern = r'\d\. ([A-Za-z]{3}/\d{4} - [A-Za-z]{3}/\d{4}|current|present)'
        experience_dates = re.findall(date_range_pattern,text)
        date_ranges = [f"{m[0]} - {m[1]}" for m in experience_dates]
        print(experience_dates)
        return experience_dates
        # print(date_ranges)
        # return date_ranges
        
        # return []  # Return an empty list if no experience dates section is found

    def calculate_months(self, start_str, end_str):
        month_map = {
            'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
            'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12,
            '01': 1, '02': 2, '03': 3, '04': 4, '05': 5, '06': 6,
            '07': 7, '08': 8, '09': 9, '10': 10, '11': 11, '12':12 
        }
        print(start_str, end_str)
        # Extract month and year from the date string
        start_str = start_str.lower()
        start_month, start_year = start_str.split('/')
        if end_str.lower() not in ["cur", "pre"]:
            end_month, end_year = end_str.split('/')
        else:
            end_month, end_year = "0"+str(datetime.now().month), datetime.now().year
        
        # Convert month abbreviation to number
        start_month_num = month_map[start_month]
        end_month_num = month_map[end_month.lower()]
        
        # Calculate the difference in years and months
        start_date = datetime(int(start_year), start_month_num, 1)
        end_date = datetime(int(end_year), end_month_num, 1)
        
        delta = (end_date.year - start_date.year) * 12 + end_date.month - start_date.month+1
        return delta

    def extract_education_experience_projects(self, text):
        """Extract education, experience, projects information from resume text"""
        LLM_response = self.extract_resume_details_using_LLM(text)
        print("\nLLM response for section: \n", LLM_response)
        sections = {}
        # Define section headers
        headers = ["Education", "Experience", "Projects"]
        
        # Split response into lines
        lines = LLM_response.split("\n")
        
        current_section = None
        for line in lines:
            line = line.strip()
            
            # Check if the line is a section header
            if any(header in line for header in headers):
                current_section = line.strip(r"#|*")
                # if "Total number of work experience" in current_section:
                #     sections[current_section] = {}  # Use dict for experience duration
                # else:
                sections[current_section] = []  # Use list for other sections
            elif current_section and line.startswith("- "):
                if isinstance(sections[current_section], list):
                    sections[current_section].append(line.strip("- "))
                else:
                    key_value = line.strip("* ").split(":", 1)
                    if len(key_value) == 2:
                        key, value = map(str.strip, key_value)
                        sections[current_section][key] = value
        # print("SECTION INFO: \n", sections)
        # section extraction completed.

        # Extracting years of experience.
        experience_dates = self.extract_experience_dates_from_section(LLM_response)
        # print("EXP DATES: \n", experience_dates)
        # Regex pattern to extract date ranges (e.g., May/2024 - May/2024)
        # date_pattern = r'({\w{3}|\d{2}}/\d{4}) - ({\w{3}|\d{2}}/\d{4})'
        date_pattern = r'(\b(?:[A-Za-z]{3}|\d{2})/\d{4}) - (\b(?:[A-Za-z]{3}|\d{2})/\d{4}|current|present)'
        # Initialize total months and years
        total_months = 0

        # Process each entry in the data
        for entry in experience_dates:
            # Find date ranges in the entry
            match = re.search(date_pattern, entry)
            # print(match)
            if match:
                start_date = match.group(1)
                end_date = match.group(2)
                total_months += self.calculate_months(start_date, end_date) 
        
        # Convert total months to years and remaining months
        total_years = total_months // 12
        remaining_months = total_months % 12
        # print(sections)
        print(f"Total Duration: {total_years} years and {remaining_months} months")

        sections["Total Experience"] = f"{total_years} years and {remaining_months} months"

        return sections

    def extract_skills(self, text):
        predefined_skills = [
            "Python", "Java", "C++", "JavaScript", "SQL", "Data Structures", 
            "Algorithms", "Git", "REST APIs", "Cloud Computing", "Docker", 
            "Kubernetes", "Machine Learning", "Deep Learning", 
            "Data Visualization", "Pandas", "NumPy", "Scikit-learn", 
            "TensorFlow", "PyTorch", "Big Data", "Hadoop", "Spark", "Linux", 
            "CI/CD", "Jenkins", "Terraform", "AWS", "Azure", "Google Cloud", 
            "Bash", "Monitoring", "Prometheus", "Grafana", "Network Security", 
            "Penetration Testing", "Ethical Hacking", "Firewalls", "SIEM", 
            "SOC Operations", "Encryption", "Threat Intelligence", 
            "Incident Response", "Kali Linux", "Metasploit", "Wireshark", 
            "CISSP", "Serverless Computing", "Networking", "React.js", 
            "Node.js", "Django", "Flask", "MongoDB", "GraphQL", "HTML", "CSS", 
            "TypeScript", "SASS", "Webpack", "UI/UX Design", "Responsive Design", 
            "Cross-Browser Testing", "Express.js", "Spring Boot", "Ruby on Rails", 
            "Redis", "Microservices", "Keras", "Computer Vision", "NLP", 
            "Data Engineering", "MLOps", "MySQL", "PostgreSQL", "Oracle DB", 
            "Database Optimization", "Indexing", "Backup & Recovery", 
            "Data Security", "ETL Pipelines", "Windows", "Troubleshooting", 
            "Technical Support", "Active Directory", "Help Desk", "VPN", 
            "Remote Desktop", "Cloud Support", "Excel", "Tableau", "Power BI", 
            "Business Intelligence", "Stakeholder Communication", 
            "Process Improvement", "Project Management"
        ]

        skill_taxonomy = {"sql": "Structured Query Language", "nlp":"Natural Language Processing",
            "css":"Cascading Style Sheet", "dl":"Deep Learning", "ml":"Machine Learning",
            "js":"JavaScript", "aws": "Amazon Web Servies", "eda":"Exploratory Data Analysis",
            "sass":"Syntactically Awesome Style Sheets", "ai":"Artificial Intelligence",
            "tf": "TensorFlow", "react":"ReactJS", "go":"Golang", "rb": "Ruby","c++": "CPP",
            "ts": "TypeScript", "node": "Node.js", "vue":"Vue.js", "express": "Express.js",
            "c#": "C-Sharp", "Azure": "Microsoft Azure"
            }

        
        skills = set()
        
        for i in predefined_skills:
            if i.lower() in text.lower():
                skills.add(i)
        skill_list = list(skills)
        # print(skill_list)
        for i in range(0,len(skill_list)):
            # print("Current: ", skill_taxonomy[] )
            if skill_list[i].lower() in skill_taxonomy.keys():
                skill_list[i] = skill_taxonomy[skill_list[i].lower()]
        
        # print(skill_list)
        return skill_list


    def extractUniqueWordsAndSentences(self, text):  # used to be extractUniqueWordsAndSentence
        """Extract unique words and sentence"""
        strText = str(text)
        uniqueWords = set(strText.split(" "))
        sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s', text)
        return sentences 

    def extract_summary(self, text):
        """Extract summary/objective from resume text"""
        summary = []
        lines = text.split('\n')
        summary_keywords = [
            'summary', 'professional summary', 'career summary', 'objective',
            'career objective', 'professional objective', 'about me', 'profile',
            'professional profile', 'career profile', 'overview', 'skill summary'
        ]
        in_summary_section = False
        current_entry = []

        # Try to find summary at the beginning of the resume
        start_index = 0
        while start_index < min(10, len(lines)) and not lines[start_index].strip():
            start_index += 1

        # Check first few non-empty lines for potential summary
        first_lines = []
        lines_checked = 0
        for line in lines[start_index:]:
            if line.strip():
                first_lines.append(line.strip())
                lines_checked += 1
                if lines_checked >= 5:  # Check first 5 non-empty lines
                    break

        # If first few lines look like a summary (no special formatting, no contact info)
        if first_lines and not any(keyword in first_lines[0].lower() for keyword in summary_keywords):
            potential_summary = ' '.join(first_lines)
            if len(potential_summary.split()) > 10:  # More than 10 words
                if not re.search(r'\b(?:email|phone|address|tel|mobile|linkedin)\b', potential_summary.lower()):
                    summary.append(potential_summary)
        # c => 189 + 140, d =  65 + 60, 270
        # Look for explicitly marked summary section
        for line in lines:
            line = line.strip()
            # Check for section header
            if any(keyword.lower() in line.lower() for keyword in summary_keywords):
                if not any(keyword.lower() == line.lower() for keyword in summary_keywords):
                    # This line contains summary info, not just a header
                    current_entry.append(line)
                in_summary_section = True
                continue
            
            if in_summary_section:
                # Check if we've hit another section
                if line and any(keyword.lower() in line.lower() for keyword in self.document_types['resume']):
                    if not any(sum_key.lower() in line.lower() for sum_key in summary_keywords):
                        in_summary_section = False
                        if current_entry:
                            summary.append(' '.join(current_entry))
                            current_entry = []
                        continue
                
                if line:
                    current_entry.append(line)
                elif current_entry:  # Empty line and we have content
                    summary.append(' '.join(current_entry))
                    current_entry = []
        
        if current_entry:
            summary.append(' '.join(current_entry))
        
        return ' '.join(summary) if summary else ''

    import requests


    def get_feedback_from_groq(self, text, job_requirements):
        """Sends the extracted text to the Groq API and gets feedback."""
        payload = {
            "model": "gemma2-9b-it",  # Change to the model you prefer
            "temperature": 0.2,
            "messages": [
                {"role": "system", "content": """YOU ARE AN AI-POWERED ATS (APPLICANT TRACKING SYSTEM) 
                 RESUME EVALUATOR, DESIGNED TO DELIVER DETAILED, PROFESSIONAL ANALYSIS OF A USER’S RESUME. 
                 YOUR PRIMARY TASK IS TO EVALUATE THE RESUME AGAINST INDUSTRY STANDARDS FOR ATS COMPATIBILITY, 
                 SCORING IT ACROSS MULTIPLE DIMENSIONS, 
                 AND PROVIDING DETAILED RECOMMENDATIONS FOR IMPROVEMENT."""},
                {"role": "user", "content": f"""
                1. PRIMARY OBJECTIVES

   - CALCULATE an Overall ATS Score (0–100), adapting to the user’s industry and job role. Format: Overall ATS Score: (0-100)/100
   - BREAK DOWN the score into six weighted categories: 
     - Keyword Optimization (25%) - Relevance, density, and placement of industry-specific terms based on job descriptions.
     - Work Experience & Achievements (20%) - Use of action verbs, quantifiable impact, and chronological structure.
     - Skills & Competencies (15%) - Balance between hard and soft skills, alignment with job expectations.
     - Education & Certifications (10%) - Completeness, correct formatting, relevance.
     - Grammar & Consistency (10%) - Uniform tense usage, spelling, punctuation, readability.
   - EXPLAIN & RECOMMEND actionable improvements tailored to the user’s job role and industry.
   - [MUST] Do not use any word formatting as this output is going to undergo futher process like data extraction [Using regular expression]. 
        Thus output should be in raw format.

2. EVALUATION PROCESS

   - READ & INTERPRET the resume text.
     - Identify ATS-blocking elements (e.g., images, graphics, tables, non-standard fonts).
     - Extract key information to assess parsing efficiency.
   - EVALUATE and assign a score to each category (out of 100%).
     - Strengths: Highlight positive aspects contributing to ATS compatibility.
     - Weaknesses: Identify formatting issues, missing keywords, inconsistencies.
   - SUPPORT WITH EVIDENCE by referencing specific resume sections where applicable.
   - RECOMMEND IMPROVEMENTS using clear, step-by-step guidance.
   - SUMMARIZE FINDINGS with:
     - Overall ATS Score
     - Category-wise Scores & Justifications
     - Consolidated Improvement Checklist (actionable next steps)

3. CHAIN OF THOUGHT LOGIC

   - UNDERSTAND ATS rules, industry-specific norms, and scoring impact.
   - BASICS of optimal ATS formatting (fonts, headings, bullet points, structured layout).
   - BREAK DOWN each category with specific criteria:
     - Keyword Optimization: Job-specific terminology, correct placements, avoiding keyword stuffing.
     - Work Experience: Actionable verbs, quantifiable achievements, structured bullets.
     - Skills & Competencies: Well-defined hard vs. soft skills, aligned with role expectations.
     - Education & Certifications: Correct ordering, proper formatting, completeness.
     - Grammar & Consistency: Checking for tense shifts, punctuation errors, readability.
   - ANALYZE against known ATS parsing behaviors (e.g., Taleo, Workday, Greenhouse compatibility).
   - BUILD a structured, expert-level report with specific suggestions and examples.
   - HANDLE EDGE CASES:
     - Career Changers: Emphasize transferable skills, relevant projects, highlight adaptability.
     - Employment Gaps: Suggest reframing sections to focus on skills, certifications, or projects.
     - Technical vs. Non-Technical Resumes: Adjust scoring to reflect domain-specific priorities.
   - FINAL ANSWER FORMAT:
     - Overall Score
     - Category Breakdown
     - Actionable Recommendations with Examples

4. ATS SYSTEM ADAPTATION

   - Ensure Resume Compatibility with ATS Parsers:
     - Test parsing against systems like Taleo, Workday, and Greenhouse.
     - Flag common parsing issues (e.g., headers in tables, missing section titles).
     - Recommend user-friendly alternatives for unsupported formatting elements.
   - Provide ATS Pre-Check Recommendations:
     - Suggest free ATS testing tools for validation.
     - Guide users on keyword density and placement best practices.

5. MODEL ADAPTATION BASED ON SIZE

   - Smaller Models (1B–7B parameters):
     - Use straightforward, highly structured instructions.
     - Provide direct feedback without excessive jargon.
     - Avoid speculative reasoning (reduce hallucinations).
   - Larger Models (13B–175B parameters):
     - Offer in-depth justifications and industry-specific insights.
     - Include advanced ATS optimization strategies (e.g., LinkedIn synergy, resume-tailoring techniques).
     - Provide role-specific resume benchmarks (e.g., software engineer vs. sales manager examples).

6. NEGATIVE PROMPTING (WHAT NOT TO DO)

   - NEVER give vague or generic feedback; insights must be pinpointed and data-driven.
   - DO NOT ignore missing keywords; suggest additions based on job descriptions.
   - AVOID unstructured suggestions; all improvement steps must be clear and actionable.
   - NEVER downplay critical issues like poor formatting, lack of structure, or improper keyword usage.
                
Here is an Example ouput:
Resume Evaluation Report for <name of the candidate>

Overall ATS Score: 82/100

Category Breakdown:

1. Keyword Optimization (25%)
- Score: 20/25
- Strength: The resume includes several relevant keywords such as Python, Pandas, Power BI, and Machine Learning.
- Weakness: Keywords like "Data Visualization," "NLP," "Deep Learning," and "AI" are mentioned but not emphasized.
- Recommendations: Add more emphasis on keywords like "Data Visualization," "NLP," and "Deep Learning."

2. *Work Experience & Achievements (20%)
- Score: 16/20
- Strength: The user has a clear and structured work history with specific roles and responsibilities.
- Weakness: The roles listed (Business Analyst, Data Extraction Engineer, Technical Recruiter) do not fully align with the Data Analyst role.
- Recomadation: Tailor the work experience to highlight data-related achievements and responsibilities.

Similary other section heading and points.
                    

            ## JOB DESCRIPTION FOR ROLE ##
                {job_requirements}.
            here is the content of resume  
                {text}."""}
            ]
        }
        # Groq API configuration
        API_KEY = "<YOUR_API_KEY>"  # Replace with your actual API key 
        URL = "https://api.groq.com/openai/v1/chat/completions"
        HEADERS = {
            "Authorization": f"Bearer {API_KEY}",
            "Content-Type": "application/json"
        }


        response = requests.post(URL, headers=HEADERS, json=payload)
        # st.write(response.json())
        response = response.json()
        return response.get("choices", [{}])[0].get("message", {}).get("content", "No response received.")


    def analyze_resume(self, resume_data, job_requirements):
        """Analyze resume and return scores and recommendations"""
        text = resume_data.get('raw_text', '')
        
        # Extract personal information
        personal_info = self.extract_personal_info(text)
        
        # First detect document type
        doc_type = self.detect_document_type(text)
        if doc_type != 'resume':
            return {
                'ats_score': 0,
                'document_type': doc_type,
                'keyword_match': {'score': 0, 'found_skills': [], 'missing_skills': []},
                'section_score': 0,
                'format_score': 0,
                'suggestions': [f"This appears to be a {doc_type} document. Please upload a resume for ATS analysis."]
            }
            
        # Calculate keyword match
        required_skills = job_requirements.get('required_skills', [])
        keyword_match = self.calculate_keyword_match(text, required_skills)
        
        # Extract all resume sections
        skills = list(self.extract_skills(text))  # Convert skills set to list
        summary = self.extract_summary(text)
        other_details = self.extract_education_experience_projects(text)
        education = other_details["Education:"]
        experience = other_details["Experience:"]
        projects = other_details["Projects:"]
        total_experience = other_details["Total Experience"]
        print(total_experience)
        # Check resume sections
        section_score = self.check_resume_sections(text)
        
        # Check formatting
        format_score, format_deductions = self.check_formatting(text)
        
        # Generate section-specific suggestions
        contact_suggestions = []
        if not personal_info.get('email'):
            contact_suggestions.append("Add your email address")
        if not personal_info.get('phone'):
            contact_suggestions.append("Add your phone number")
        if not personal_info.get('linkedin'):
            contact_suggestions.append("Add your LinkedIn profile URL")
        
        summary_suggestions = []
        if not summary:
            summary_suggestions.append("Add a professional summary to highlight your key qualifications")
        elif len(summary.split()) < 30:
            summary_suggestions.append("Expand your professional summary to better highlight your experience and goals")
        elif len(summary.split()) > 100:
            summary_suggestions.append("Consider making your summary more concise (aim for 50-75 words)")
        
        skills_suggestions = []
        if not skills:
            skills_suggestions.append("Add a dedicated skills section")
        if isinstance(skills, (list, set)) and len(list(skills)) < 5:
            skills_suggestions.append("List more relevant technical and soft skills")
        if keyword_match['score'] < 70:
            skills_suggestions.append("Add more skills that match the job requirements")
        
        experience_suggestions = []
        if not experience:
            experience_suggestions.append("Add your work experience section")
        else:
            has_dates = any(re.search(r'\b(19|20)\d{2}\b', exp) for exp in experience)
            has_bullets = any(re.search(r'[•\-\*]', exp) for exp in experience)
            has_action_verbs = any(re.search(r'\b(developed|managed|created|implemented|designed|led|improved)\b', 
                                           exp.lower()) for exp in experience)
            
            if not has_dates:
                experience_suggestions.append("Include dates for each work experience")
            if not has_bullets:
                experience_suggestions.append("Use bullet points to list your achievements and responsibilities")
            if not has_action_verbs:
                experience_suggestions.append("Start bullet points with strong action verbs")
        
        education_suggestions = []
        if not education:
            education_suggestions.append("Add your educational background")
        else:
            has_dates = any(re.search(r'\b(19|20)\d{2}\b', edu) for edu in education)
            has_degree = any(re.search(r'\b(bachelor|master|phd|b\.|m\.|diploma)\b', 
                                     edu.lower()) for edu in education)
            has_gpa = any(re.search(r'\b(gpa|cgpa|grade|percentage)\b', 
                                  edu.lower()) for edu in education)
            
            if not has_dates:
                education_suggestions.append("Include graduation dates")
            if not has_degree:
                education_suggestions.append("Specify your degree type")
            if not has_gpa and job_requirements.get('require_gpa', False):
                education_suggestions.append("Include your GPA if it's above 3.0")
        
        format_suggestions = []
        if format_score < 100:
            format_suggestions.extend(format_deductions)
        
        # Calculate section-specific scores
        contact_score = 100 - (len(contact_suggestions) * 25)  # -25 for each missing item
        summary_score = 100 - (len(summary_suggestions) * 33)  # -33 for each issue
        skills_score = keyword_match['score']
        experience_score = 100 - (len(experience_suggestions) * 25)
        education_score = 100 - (len(education_suggestions) * 25)
        
        # Calculate overall ATS score with weighted components
        LLM_feedback = self.get_feedback_from_groq(text, job_requirements) 
        print("LLM FEEDBACK: ",LLM_feedback)


        pattern = r'ATS Score:.* (\d\d\/100)'
        # Find matches using regex
        matches = re.findall(pattern, LLM_feedback, re.IGNORECASE)
        print(matches)
        ats_score = int(matches[0][0:2])
        
        print("ATS SCORE: ",ats_score)

        pattern_summary = re.findall(r"(\d+\..*?)\n- Score:.*?\n- Strength: (.*?)\n- Weakness: (.*?)\n", LLM_feedback, re.DOTALL)

        # matchs in form of dictonary
        # matches = [match.groupdict() for match in pattern_summary.finditer(LLM_feedback)] 
        summary_feedback = pattern_summary
        summary_feedback.append(ats_score)
        for section in summary_feedback:
            print(section)

        # ats_score = (
        #     int(round(contact_score * 0.1)) +      # 10% weight for contact info
        #     int(round(summary_score * 0.1)) +      # 10% weight for summary
        #     int(round(skills_score * 0.3)) +       # 30% weight for skills match
        #     int(round(experience_score * 0.2)) +   # 20% weight for experience
        #     int(round(education_score * 0.1)) +    # 10% weight for education
        #     int(round(format_score * 0.2))         # 20% weight for formatting
        # )
        
        # Combine all suggestions into a single list
        suggestions = []
        suggestions.extend(contact_suggestions)
        suggestions.extend(summary_suggestions)
        suggestions.extend(skills_suggestions)
        suggestions.extend(experience_suggestions)
        suggestions.extend(education_suggestions)
        suggestions.extend(format_suggestions)
        
        if not suggestions:
            suggestions.append("Your resume is well-optimized for ATS systems")
        
        return {
            **personal_info,  # Include extracted personal info
            'ats_score': ats_score,
            'document_type': 'resume',
            'keyword_match': keyword_match,
            'section_score': section_score,
            'format_score': format_score,
            'education': education,
            'experience': experience,
            'total_experience': total_experience,
            'projects': projects,
            'skills': skills,
            'summary': summary,
            'suggestions': suggestions,
            'contact_suggestions': contact_suggestions,
            'summary_suggestions': summary_suggestions,
            'skills_suggestions': skills_suggestions,
            'experience_suggestions': experience_suggestions,
            'education_suggestions': education_suggestions,
            'format_suggestions': format_suggestions,
            'summary_feedback': summary_feedback,
            'section_scores': {
                'contact': contact_score,
                'summary': summary_score,
                'skills': skills_score,
                'experience': experience_score,
                'education': education_score,
                'format': format_score
            }
        }
    
    def extract_links_from_pdf(self,pdf_file):
        links = []
        reader = PdfReader(pdf_file)
        for page in reader.pages:
            if "/Annots" in page:  # Check if annotations exist
                for annot in page["/Annots"]:
                    annot_obj = annot.get_object()  # Get annotation object
                    if "/A" in annot_obj:  # Check if action dictionary exists
                        action = annot_obj["/A"]
                        if "/URI" in action:  # Check if URI exists
                            links.append(action["/URI"])  # Extract link
        return links    

    def extract_links_from_docx(self,docx_file):
        doc = Document(docx_file)
        links = []
        for rel in doc.part.rels:
            if "hyperlink" in doc.part.rels[rel].reltype:
                links.append(doc.part.rels[rel].target_ref)
        return links

    def check_url_status(self,url):
        try:
            response = requests.head(url, allow_redirects=True, timeout=5)
            return response.status_code == 200
        except requests.RequestException:
            return False
        
    def to_excel(self, df):
        output = BytesIO()
        with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
            df.to_excel(writer, index=False, sheet_name="Sheet1")
        return output.getvalue()
    
